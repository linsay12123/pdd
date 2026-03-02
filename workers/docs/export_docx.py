import json
import sys
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor


def apply_font(run):
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(0, 0, 0)


def set_hanging_indent(paragraph):
    paragraph.paragraph_format.left_indent = Pt(24)
    paragraph.paragraph_format.first_line_indent = Pt(-24)


def set_spacing(paragraph):
    paragraph.paragraph_format.space_after = Pt(12)


def add_blank_paragraph(document):
    paragraph = document.add_paragraph("")
    set_spacing(paragraph)


def build_docx(payload, output_path):
    document = Document()

    title_paragraph = document.add_paragraph()
    title_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title_paragraph.add_run(payload["title"])
    title_run.bold = True
    apply_font(title_run)
    title_run.font.size = Pt(14)
    set_spacing(title_paragraph)

    add_blank_paragraph(document)

    for section in payload["sections"]:
        heading_paragraph = document.add_paragraph()
        heading_run = heading_paragraph.add_run(section["heading"])
        heading_run.bold = True
        apply_font(heading_run)
        set_spacing(heading_paragraph)

        add_blank_paragraph(document)

        for text in section["paragraphs"]:
            paragraph = document.add_paragraph()
            run = paragraph.add_run(text)
            apply_font(run)
            set_spacing(paragraph)
            add_blank_paragraph(document)

    references_heading = document.add_paragraph()
    references_run = references_heading.add_run("References")
    references_run.bold = True
    apply_font(references_run)
    set_spacing(references_heading)

    add_blank_paragraph(document)

    for reference in payload["references"]:
        paragraph = document.add_paragraph()
        run = paragraph.add_run(reference)
        apply_font(run)
        set_hanging_indent(paragraph)
        set_spacing(paragraph)

    output_path.parent.mkdir(parents=True, exist_ok=True)
    document.save(output_path)


def build_sample_payload():
    return {
        "taskId": "sample",
        "title": "Sample Final Article",
        "sections": [
            {
                "heading": "Introduction",
                "paragraphs": [
                    "This sample paragraph confirms the DOCX export pipeline is working."
                ],
            }
        ],
        "references": ["Smith, A. (2024). Sample source."],
        "citationStyle": "APA 7",
    }


def main():
    if len(sys.argv) == 2 and sys.argv[1] == "--sample":
        output_path = Path("output/doc/sample-final.docx")
        build_docx(build_sample_payload(), output_path)
        print(output_path)
        return

    if len(sys.argv) != 3:
        raise SystemExit("Usage: export_docx.py <payload.json> <output.docx>")

    payload_path = Path(sys.argv[1])
    output_path = Path(sys.argv[2])
    payload = json.loads(payload_path.read_text(encoding="utf-8"))
    build_docx(payload, output_path)
    print(output_path)


if __name__ == "__main__":
    main()
